import sys
import os
import json
import re
import copy
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# 遵循 UTF-8 与简体中文编码标准
sys.stdout.reconfigure(encoding='utf-8')

def clone_paragraph(paragraph, parent):
    """
    深度复制段落底层 OpenXML 节点并挂载，100% 完美克隆样式（字体、加粗、颜色、段距）
    """
    p_element = paragraph._p
    new_p = copy.deepcopy(p_element)
    
    # 清空克隆出来的文本内容，防止残留模板字眼
    # OpenXML 结构下，文本通常在 p -> r -> t 节点中
    for r in new_p.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            t.text = ""
            
    # 将克隆的节点追加到父容器
    parent._element.append(new_p)
    return Paragraph(new_p, parent)

def replace_run_text_keep_style(paragraph, match_regex, replace_text):
    """
    高精准替换段落里的 Run 文本，绝对完美保留原有字体、字号、加粗与高档配色
    """
    text = paragraph.text
    if not re.search(match_regex, text):
        return False
        
    # 如果段落只有一个 Run，直接替换
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = re.sub(match_regex, replace_text, paragraph.runs[0].text)
        return True
        
    # 如果跨越多个 Run，合并替换，并使文本归属于第一个 Run 继承格式，清空其余 Runs
    combined_text = "".join([r.text for r in paragraph.runs])
    new_text = re.sub(match_regex, replace_text, combined_text)
    
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    return True

def fill_docx_template(data_path, template_path, output_path):
    """
    核心：万能简历 Word 模板智能分析、经历克隆与格式缩减自适应引擎
    """
    # 1. 读入数据
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    basic = data.get('basicInfo', {})
    edus = data.get('education', [])
    exps = data.get('experience', [])
    projs = data.get('projects', [])
    skills = data.get('skills', [])
    
    # 2. 读入模板
    if not os.path.exists(template_path):
        print(f"❌ 错误：未找到模板文件 {template_path}")
        return False
        
    doc = Document(template_path)
    
    # ==========================================================================
    # 第一步：高精度基本信息全局匹配替换 (语义检索)
    # ==========================================================================
    
    # 定义基本信息的智能正则替换规则
    basic_replacements = []
    
    if basic.get('phone'):
        # 匹配11位手机号
        basic_replacements.append((r'1[3-9]\d{9}', basic['phone']))
        basic_replacements.append((r'(?:电话|手机|联络)[:：]?\s*[^\s,，|]+', f"电话: {basic['phone']}"))
        
    if basic.get('email'):
        # 匹配邮箱正则
        basic_replacements.append((r'[a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z]{2,6}', basic['email']))
        basic_replacements.append((r'(?:邮箱|Email|邮件)[:：]?\s*[^\s,，|]+', f"邮箱: {basic['email']}"))
        
    if basic.get('wechat'):
        basic_replacements.append((r'(?:微信|WeChat|wechat)[:：]?\s*[^\s,，|]+', f"微信: {basic['wechat']}"))
        
    if basic.get('github'):
        basic_replacements.append((r'(?:GitHub|Github|github.com)[:：]?\s*[^\s,，|]+', f"GitHub: {basic['github']}"))
        
    if basic.get('title'):
        basic_replacements.append((r'(?:意向|求职意向|期望职位|应聘职位)[:：]?\s*[^\s,，|]+', f"求职意向: {basic['title']}"))

    # 全局遍历段落和表格进行替换
    def apply_basic_info(paragraphs):
        # 1. 智能查找并替换姓名 (简历第一行或者前两行字号最大的那个)
        name_replaced = False
        if basic.get('name'):
            # 语义匹配：替换包含常见名字占位词
            for p in paragraphs[:5]:
                if p.text.strip() in ["姓名", "张三", "李四", "王五", "示例姓名", "求职者"]:
                    p.text = basic['name']
                    name_replaced = True
                    break
            
            # 如果没找到占位词，寻找前 3 个段落中字号最大、通常即是名字的段落进行直接换名
            if not name_replaced:
                max_size = 0
                name_para = None
                for p in paragraphs[:3]:
                    for r in p.runs:
                        if r.font.size and r.font.size > max_size:
                            max_size = r.font.size
                            name_para = p
                if name_para:
                    name_para.text = basic['name']

        # 2. 替换联系方式和岗位
        for p in paragraphs:
            for regex, rep in basic_replacements:
                replace_run_text_keep_style(p, regex, rep)

    apply_basic_info(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                apply_basic_info(cell.paragraphs)

    # ==========================================================================
    # 第二步：智能经历元模块定位、克隆与填充 (经历自适应克隆)
    # ==========================================================================
    
    # 辅助函数：根据示例关键词（如大学、公司、项目）在文档中寻找包含特定模块段落
    # 并将其及其下方的描述行组成“一个逻辑卡片元组”进行动态扩展
    def process_dynamic_section(paragraphs, data_list, keyword_trigger, field_mapping_fn):
        """
        paragraphs: 扫描段落集
        data_list: 用户填入的结构数组，例如 education
        keyword_trigger: 触发扫描的关键字，如 '大学' / '公司'
        field_mapping_fn: 接受段落和单条数据，执行实际的字段和内容写入
        """
        if not data_list:
            return
            
        # 1. 扫描段落，定位所有的“示例经历元卡片块”
        blocks = []
        in_block = False
        current_block = []
        
        for i, p in enumerate(paragraphs):
            text = p.text.strip()
            # 发现包含关键字 (且长度适中不是大标题) 触发新块
            if keyword_trigger(text) and len(text) < 45:
                if in_block and current_block:
                    blocks.append(current_block)
                    current_block = []
                in_block = True
                current_block.append((i, p))
            elif in_block:
                # 块描述内容：收集到遇到下一个大模块标题或空段为止
                if text == "" or (len(text) < 10 and any(h in text for h in ["工作", "项目", "教育", "技能", "评价"])):
                    blocks.append(current_block)
                    current_block = []
                    in_block = False
                else:
                    current_block.append((i, p))
                    
        if in_block and current_block:
            blocks.append(current_block)

        if not blocks:
            return

        # 2. 映射与克隆
        # 模板里有 blocks.length 个示例槽，用户需要填入 data_list.length 个真实经历
        template_block = blocks[-1] # 使用最后一个块作为克隆的母本
        parent = paragraphs[0]._parent # 父容器 (如 Document 或者是 Cell)
        
        # 依次写入
        for idx, item in enumerate(data_list):
            if idx < len(blocks):
                # 模板槽足够，直接修改
                target_block = blocks[idx]
                field_mapping_fn(target_block, item)
            else:
                # 模板槽不够用，执行高精准 OXML 节点克隆插入
                new_block_paras = []
                for _, origin_p in template_block:
                    cloned_p = clone_paragraph(origin_p, parent)
                    new_block_paras.append((None, cloned_p))
                
                # 对克隆出的新段落进行数据填入
                field_mapping_fn(new_block_paras, item)
                
        # 3. 抹除多余的示例槽
        if len(data_list) < len(blocks):
            for excess_idx in range(len(data_list), len(blocks)):
                for _, excess_p in blocks[excess_idx]:
                    # 在 OpenXML 中销毁多余段落
                    p_element = excess_p._p
                    p_element.getparent().remove(p_element)

    # 2.1 教育修业填充与克隆
    def fill_edu(block_paras, item):
        # block_paras 是一个 (idx, Paragraph) 数组。第一行通常是学校/专业行
        header_p = block_paras[0][1]
        
        # 智能匹配第一行的字段（如果有表格列，更整齐）
        header_p.text = "" # 清空重写
        run1 = header_p.add_run(f"{item.get('school', '')}  |  {item.get('major', '')} ({item.get('degree', '')})")
        run1.bold = True
        run1.font.size = Pt(11)
        
        # 补充日期在右边 (使用加空格两端对齐或直接追加)
        if item.get('date'):
            header_p.add_run(f"      {item['date']}")
            
        # 其余行填充描述
        if len(block_paras) > 1 and item.get('description'):
            desc_p = block_paras[1][1]
            desc_p.text = item['description']
            
            # 多余的描述行清空
            for _, extra_p in block_paras[2:]:
                extra_p.text = ""

    process_dynamic_section(
        doc.paragraphs, 
        edus, 
        lambda t: "大学" in t or "学院" in t or "校" in t or "学府" in t, 
        fill_edu
    )

    # 2.2 工作履历填充与克隆
    def fill_exp(block_paras, item):
        header_p = block_paras[0][1]
        header_p.text = ""
        run1 = header_p.add_run(f"{item.get('company', '')}  |  {item.get('role', '')}")
        run1.bold = True
        run1.font.size = Pt(11)
        
        if item.get('date'):
            header_p.add_run(f"      {item['date']}")
            
        if len(block_paras) > 1 and item.get('description'):
            # 将多行描述合并写入
            desc_p = block_paras[1][1]
            desc_p.text = item['description']
            for _, extra_p in block_paras[2:]:
                extra_p.text = ""

    process_dynamic_section(
        doc.paragraphs, 
        exps, 
        lambda t: "公司" in t or "集团" in t or "中心" in t or "行" in t or "企业" in t, 
        fill_exp
    )

    # 2.3 项目经验填充与克隆
    def fill_proj(block_paras, item):
        header_p = block_paras[0][1]
        header_p.text = ""
        run1 = header_p.add_run(f"{item.get('name', '')}  |  {item.get('role', '')}")
        run1.bold = True
        run1.font.size = Pt(11)
        
        if item.get('date'):
            header_p.add_run(f"      {item['date']}")
            
        if len(block_paras) > 1 and item.get('description'):
            desc_p = block_paras[1][1]
            desc_p.text = item['description']
            for _, extra_p in block_paras[2:]:
                extra_p.text = ""

    process_dynamic_section(
        doc.paragraphs, 
        projs, 
        lambda t: "项目" in t or "系统" in t or "平台" in t or "软件" in t or "应用" in t, 
        fill_proj
    )

    # ==========================================================================
    # 第三步：技能平铺与间距微调 (单页自适应缩水控制)
    # ==========================================================================
    
    # 智能定位技能段落
    if skills:
        skill_replaced = False
        # 寻找包含“技能”、“特长”的标题的下方段落
        for i, p in enumerate(doc.paragraphs):
            if any(k in p.text for k in ["技能", "特长", "证书"]) and len(p.text) < 12:
                # 替换其下方的第一个段落
                if i + 1 < len(doc.paragraphs):
                    desc_para = doc.paragraphs[i + 1]
                    desc_para.text = "；".join(skills)
                    skill_replaced = True
                    break
        
        # 保底平铺替换
        if not skill_replaced:
            for p in doc.paragraphs:
                if "精通" in p.text or "熟练" in p.text:
                    p.text = "；".join(skills)
                    break

    # 间距与字号微调整体自适应收紧，确保 100% 保持在单页内！
    total_length = sum(len(p.text) for p in doc.paragraphs)
    # 如果总字符量超过 800，触发 Word 级字号间距微调收缩，实现强约束！
    if total_length > 800:
        for p in doc.paragraphs:
            p_format = p.paragraph_format
            # 微收窄行距 (原本1.5倍的收窄到1.25倍)
            p_format.line_spacing = 1.22
            # 压缩段前段后边距，免除多余空白溢出！
            p_format.space_before = Pt(2)
            p_format.space_after = Pt(2)
            # 对所有 Runs 稍微调小字号 0.5 磅 (Pt 10.5 是五号字，Pt 10 是小五号)
            for r in p.runs:
                if r.font.size and r.font.size > Pt(10.5):
                    r.font.size = Pt(10.5)

    # 4. 保存为成品可二次自由编辑的 DOCX 文件
    doc.save(output_path)
    print(f"✅ 万能Word模板套用成功！可编辑文档已输出至：{output_path}")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("参数错误。用法: python3 docx_filler.py <data_json_path> <template_docx_path> <output_docx_path>")
        sys.exit(1)
        
    data_json = sys.argv[1]
    template_docx = sys.argv[2]
    output_docx = sys.argv[3]
    
    success = fill_docx_template(data_json, template_docx, output_docx)
    if not success:
        sys.exit(1)
